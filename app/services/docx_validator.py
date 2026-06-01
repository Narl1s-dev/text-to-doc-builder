from dataclasses import dataclass, field
from pathlib import Path
import re
from typing import Any

from docx import Document

from app.schemas.document_spec import DocumentSpec


@dataclass(frozen=True)
class DocxValidationResult:
    valid: bool
    errors: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    text_length: int = 0
    paragraph_count: int = 0
    table_count: int = 0

    def model_dump(self) -> dict[str, Any]:
        return {
            "valid": self.valid,
            "errors": self.errors,
            "warnings": self.warnings,
            "text_length": self.text_length,
            "paragraph_count": self.paragraph_count,
            "table_count": self.table_count,
        }


class DocxValidator:
    def validate(self, file_path: Path, document_spec: DocumentSpec) -> DocxValidationResult:
        errors: list[str] = []
        warnings: list[str] = []
        file_path = file_path.resolve()

        if not file_path.exists():
            return DocxValidationResult(
                valid=False,
                errors=[f"DOCX file does not exist: {file_path}"],
            )

        if file_path.stat().st_size == 0:
            return DocxValidationResult(
                valid=False,
                errors=[f"DOCX file is empty: {file_path}"],
            )

        try:
            document = Document(file_path)
        except Exception as exc:
            return DocxValidationResult(
                valid=False,
                errors=[f"DOCX file cannot be opened by python-docx: {exc}"],
            )

        text_parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        text_parts.extend(_table_text(document))
        full_text = "\n".join(text_parts).strip()

        if not full_text:
            errors.append("DOCX file contains no readable text.")

        required_heading = _first_markdown_heading(document_spec.content_markdown)
        if required_heading and _normalize_text(required_heading) not in _normalize_text(full_text):
            errors.append(f"DOCX file is missing required heading from DocumentSpec: {required_heading}")

        if len(full_text) < 20:
            warnings.append("DOCX file contains very little text.")

        return DocxValidationResult(
            valid=not errors,
            errors=errors,
            warnings=warnings,
            text_length=len(full_text),
            paragraph_count=len(document.paragraphs),
            table_count=len(document.tables),
        )


def _table_text(document) -> list[str]:
    values: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    values.append(text)
    return values


def _first_markdown_heading(markdown: str) -> str | None:
    for line in markdown.splitlines():
        match = re.match(r"^\s{0,3}#{1,6}\s+(.+?)\s*$", line)
        if match:
            heading = match.group(1).strip(" #")
            return heading or None
    return None


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip().casefold()
