from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Pt

from app.core.config import Settings, get_settings
from app.renderers.base import RenderedArtifact
from app.schemas.artifact_plan import ArtifactBlock, ArtifactPlan
from app.schemas.document_spec import DocumentSpec
from app.schemas.formats import ArtifactFormat


class DocxRenderer:
    output_format = ArtifactFormat.docx

    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings or get_settings()

    def render(self, artifact_id: str, plan: ArtifactPlan | DocumentSpec) -> RenderedArtifact:
        storage_path = self.settings.artifact_storage_path
        storage_path.mkdir(parents=True, exist_ok=True)

        file_name = f"{artifact_id}.{self.output_format.value}"
        file_path = (storage_path / file_name).resolve()

        document = Document()
        self._apply_formatting(document, plan)
        if isinstance(plan, DocumentSpec):
            self._render_document_spec(document, plan)
        else:
            self._render_blocks(document, plan)
        document.save(file_path)

        return RenderedArtifact(
            file_name=file_name,
            file_path=file_path,
            file_size=file_path.stat().st_size,
        )

    def _apply_formatting(self, document: Document, plan: ArtifactPlan) -> None:
        formatting = plan.formatting
        section = document.sections[0]
        section.top_margin = Cm(formatting.margins.top_cm)
        section.right_margin = Cm(formatting.margins.right_cm)
        section.bottom_margin = Cm(formatting.margins.bottom_cm)
        section.left_margin = Cm(formatting.margins.left_cm)

        if formatting.page_size.upper() == "A4":
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)

        if formatting.orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

        normal_style = document.styles["Normal"]
        normal_style.font.name = formatting.font_family
        normal_style.font.size = Pt(formatting.font_size)
        normal_style.paragraph_format.line_spacing = formatting.line_spacing
        normal_style.paragraph_format.space_after = Pt(formatting.paragraph_spacing_after)

    def _render_blocks(self, document: Document, plan: ArtifactPlan) -> None:
        blocks = plan.blocks or [ArtifactBlock(type="heading", level=1, text=plan.title)]
        for block in blocks:
            match block.type:
                case "heading":
                    document.add_heading(block.text or plan.title, level=block.level or 1)
                case "paragraph":
                    document.add_paragraph(block.text or "")
                case "bullet_list":
                    for item in block.items:
                        document.add_paragraph(item, style="List Bullet")
                case "numbered_list":
                    for item in block.items:
                        document.add_paragraph(item, style="List Number")
                case "table":
                    self._add_table(document, block)

    def _add_table(self, document: Document, block: ArtifactBlock) -> None:
        rows = block.rows
        if not rows:
            return

        column_count = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=column_count)
        table.style = "Table Grid"

        for row_index, row in enumerate(rows):
            for column_index in range(column_count):
                value = row[column_index] if column_index < len(row) else ""
                table.rows[row_index].cells[column_index].text = value

    def _render_document_spec(self, document: Document, spec: DocumentSpec) -> None:
        lines = spec.content_markdown.splitlines()
        paragraph_buffer: list[str] = []

        def flush_paragraph() -> None:
            if paragraph_buffer:
                document.add_paragraph(" ".join(paragraph_buffer).strip())
                paragraph_buffer.clear()

        index = 0
        while index < len(lines):
            raw_line = lines[index]
            line = raw_line.strip()

            if not line:
                flush_paragraph()
                index += 1
                continue

            if line.startswith("|"):
                flush_paragraph()
                table_lines: list[str] = []
                while index < len(lines) and lines[index].strip().startswith("|"):
                    table_lines.append(lines[index].strip())
                    index += 1
                self._add_markdown_table(document, table_lines)
                continue

            heading_level = _markdown_heading_level(line)
            if heading_level is not None:
                flush_paragraph()
                text = line[heading_level + 1 :].strip()
                document.add_heading(text or spec.title, level=min(heading_level, 6))
                index += 1
                continue

            bullet_text = _strip_bullet_marker(line)
            if bullet_text is not None:
                flush_paragraph()
                document.add_paragraph(bullet_text, style="List Bullet")
                index += 1
                continue

            numbered_text = _strip_numbered_marker(line)
            if numbered_text is not None:
                flush_paragraph()
                document.add_paragraph(numbered_text, style="List Number")
                index += 1
                continue

            paragraph_buffer.append(line)
            index += 1

        flush_paragraph()

    def _add_markdown_table(self, document: Document, table_lines: list[str]) -> None:
        rows = [_split_markdown_table_row(line) for line in table_lines]
        rows = [row for row in rows if row and not _is_markdown_separator_row(row)]
        if not rows:
            return

        block = ArtifactBlock(type="table", rows=rows)
        self._add_table(document, block)


def _markdown_heading_level(line: str) -> int | None:
    marker_length = len(line) - len(line.lstrip("#"))
    if marker_length == 0 or marker_length > 6:
        return None
    if len(line) <= marker_length or line[marker_length] != " ":
        return None
    return marker_length


def _strip_bullet_marker(line: str) -> str | None:
    for marker in ("- ", "* "):
        if line.startswith(marker):
            text = line[len(marker) :].strip()
            return text or None
    return None


def _strip_numbered_marker(line: str) -> str | None:
    marker_index = line.find(". ")
    if marker_index <= 0:
        return None
    if not line[:marker_index].isdigit():
        return None
    text = line[marker_index + 2 :].strip()
    return text or None


def _split_markdown_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_markdown_separator_row(row: list[str]) -> bool:
    return all(cell and set(cell) <= {"-", ":"} for cell in row)
