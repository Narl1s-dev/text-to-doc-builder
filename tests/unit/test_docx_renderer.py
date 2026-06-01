from docx import Document

from app.core.config import Settings
from app.renderers.docx_renderer import DocxRenderer
from app.schemas.document_spec import DocumentSpec


def test_docx_renderer_builds_docx_from_document_spec(tmp_path) -> None:
    renderer = DocxRenderer(Settings(artifact_storage_path=tmp_path))
    document_spec = DocumentSpec(
        title="Итоги встречи",
        content_markdown="""
# Итоги встречи

Обсудили план запуска.

- Подготовить документ
- Согласовать сроки

| Роль | Ответственный |
| --- | --- |
| Документ | Анна |
""",
    )

    rendered = renderer.render("art_123", document_spec)

    document = Document(rendered.file_path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    assert rendered.file_name == "art_123.docx"
    assert "Итоги встречи" in paragraphs
    assert "Обсудили план запуска." in paragraphs
    assert "Подготовить документ" in paragraphs
    assert document.tables[0].rows[1].cells[1].text == "Анна"
