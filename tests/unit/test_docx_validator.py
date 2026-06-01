from docx import Document

from app.schemas.document_spec import DocumentSpec
from app.services.docx_validator import DocxValidator


def test_docx_validator_accepts_valid_docx(tmp_path) -> None:
    file_path = tmp_path / "valid.docx"
    document = Document()
    document.add_heading("Spec", level=1)
    document.add_paragraph("The generated document contains useful body text.")
    document.save(file_path)

    result = DocxValidator().validate(
        file_path,
        DocumentSpec(title="Spec", content_markdown="# Spec\n\nBody"),
    )

    assert result.valid is True
    assert result.errors == []
    assert result.text_length > 20


def test_docx_validator_rejects_empty_docx(tmp_path) -> None:
    file_path = tmp_path / "empty.docx"
    Document().save(file_path)

    result = DocxValidator().validate(
        file_path,
        DocumentSpec(title="Spec", content_markdown="# Spec\n\nBody"),
    )

    assert result.valid is False
    assert "no readable text" in result.errors[0]


def test_docx_validator_rejects_missing_required_heading(tmp_path) -> None:
    file_path = tmp_path / "missing-heading.docx"
    document = Document()
    document.add_paragraph("The body is present, but the heading is not.")
    document.save(file_path)

    result = DocxValidator().validate(
        file_path,
        DocumentSpec(title="Spec", content_markdown="# Spec\n\nBody"),
    )

    assert result.valid is False
    assert any("missing required heading" in error for error in result.errors)
