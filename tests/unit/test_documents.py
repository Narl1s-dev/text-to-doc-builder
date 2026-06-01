import json
import time

import pytest
from fastapi.testclient import TestClient
from docx import Document

from app.core.config import get_settings
from app.db.session import get_engine
from app.main import create_app


def wait_for_document(client: TestClient, document_id: str, timeout_seconds: float = 5.0) -> dict:
    deadline = time.monotonic() + timeout_seconds
    last_data = None

    while time.monotonic() < deadline:
        response = client.get(f"/documents/{document_id}")
        assert response.status_code == 200
        last_data = response.json()
        if last_data["status"] in {"ready", "failed", "canceled"}:
            return last_data
        time.sleep(0.05)

    pytest.fail(f"Document {document_id} did not finish in time. Last response: {last_data}")


@pytest.fixture(autouse=True)
def clear_settings_cache(monkeypatch: pytest.MonkeyPatch, tmp_path) -> None:
    monkeypatch.setenv("DATABASE_URL", "sqlite:///:memory:")
    monkeypatch.setenv("ARTIFACT_STORAGE_PATH", str(tmp_path))
    monkeypatch.setenv("OPENROUTER_API_KEY", "")
    monkeypatch.delenv("API_INTERNAL_TOKEN", raising=False)
    get_settings.cache_clear()
    get_engine.cache_clear()
    yield
    get_settings.cache_clear()
    get_engine.cache_clear()


def test_create_document_creates_ready_docx_artifact(tmp_path) -> None:
    with TestClient(create_app()) as client:
        response = client.post(
            "/documents",
            json={
                "prompt": "Сделай краткий документ по итогам встречи",
                "overrides": {"title": "Итоги встречи"},
                "metadata": {"client": "test"},
            },
        )

        queued_data = response.json()
        data = wait_for_document(client, queued_data["document_id"])

    assert response.status_code == 202
    assert queued_data["status"] == "queued"
    assert queued_data["status_url"] == f"/documents/{queued_data['document_id']}"
    assert data["document_id"].startswith("art_")
    assert data["request_id"].startswith("req_")
    assert data["job_id"].startswith("job_")
    assert data["status"] == "ready"
    assert data["current_stage"] == "completed"
    assert data["output_format"] == "docx"
    assert data["file_name"] == f"{data['document_id']}.docx"
    assert data["download_url"] == f"/documents/{data['document_id']}/download"
    assert data["error_message"] is None
    assert "OPENROUTER_API_KEY is not configured." in data["warnings"]
    assert "LLM planning was skipped; fallback defaults were used." in data["warnings"]

    document = Document(tmp_path / data["file_name"])
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    assert "Итоги встречи" in paragraphs

    document_spec_path = tmp_path / f"{data['document_id']}.document_spec.json"
    document_spec = json.loads(document_spec_path.read_text(encoding="utf-8"))
    assert document_spec["schema_version"] == "document_spec.v1"
    assert document_spec["output_format"] == "docx"
    assert "Итоги встречи" in document_spec["content_markdown"]


def test_create_document_file_can_be_downloaded() -> None:
    with TestClient(create_app()) as client:
        create_response = client.post(
            "/documents",
            json={"prompt": "Сделай краткий документ по итогам встречи"},
        )
        document_id = create_response.json()["document_id"]
        wait_for_document(client, document_id)

        info_response = client.get(f"/documents/{document_id}")
        download_response = client.get(f"/documents/{document_id}/download")

    assert info_response.status_code == 200
    assert info_response.json()["status"] == "ready"
    assert download_response.status_code == 200
    assert download_response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert download_response.content.startswith(b"PK")


def test_create_document_docx_contains_fallback_prompt(tmp_path) -> None:
    with TestClient(create_app()) as client:
        response = client.post(
            "/documents",
            json={"prompt": "Сделай краткий документ по итогам встречи"},
        )
        data = wait_for_document(client, response.json()["document_id"])

    document_path = tmp_path / data["file_name"]
    document = Document(document_path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    assert "Документ" in paragraphs
    assert "Сделай краткий документ по итогам встречи" in paragraphs


def test_create_document_rejects_empty_prompt() -> None:
    with TestClient(create_app()) as client:
        response = client.post("/documents", json={"prompt": ""})

    assert response.status_code == 422


def test_create_document_rejects_future_format_until_renderer_exists() -> None:
    with TestClient(create_app()) as client:
        response = client.post(
            "/documents",
            json={
                "prompt": "Сделай презентацию",
                "output_format": "pptx",
            },
        )

    assert response.status_code == 422
    assert "not supported yet" in response.text


def test_create_document_rejects_inferred_future_format_before_enqueue() -> None:
    with TestClient(create_app()) as client:
        response = client.post(
            "/documents",
            json={
                "prompt": "Сделай презентацию о выцветании кораллов",
            },
        )

    assert response.status_code == 422
    assert "was inferred from the prompt" in response.text
    assert "not supported yet" in response.text


def test_create_document_checks_internal_token(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("API_INTERNAL_TOKEN", "secret-token")
    get_settings.cache_clear()
    get_engine.cache_clear()

    with TestClient(create_app()) as client:
        missing_token_response = client.post("/documents", json={"prompt": "Создай документ"})
        valid_token_response = client.post(
            "/documents",
            json={"prompt": "Создай документ"},
            headers={"X-API-Token": "secret-token"},
        )

    assert missing_token_response.status_code == 401
    assert valid_token_response.status_code == 202
